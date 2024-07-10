import docx

def create_invitation(guest_name):
    doc = docx.Document()
    
    # Add a title
    doc.add_heading('Invitation', level=1)
    
    # Add invitation message
    message = (
        f"Dear {guest_name},\n\n"
        "You are cordially invited to our special event. We look forward to your presence.\n\n"
        "Sincerely,\n"
        "Event Organizer"
    )
    doc.add_paragraph(message)
    
    # Save the document
    filename = f'Invitation_{guest_name.replace(" ", "_")}.docx'
    doc.save(filename)
    print(f'Invitation created for {guest_name}: {filename}')

def create_invitations_for_guests(guests):
    for guest in guests:
        create_invitation(guest)

if __name__ == "__main__":
    # Example list of guests
    guests = [
        "Alice Smith",
        "Bob Johnson",
        "Charlie Brown",
        "Diana Ross"
    ]
    
    create_invitations_for_guests(guests)
